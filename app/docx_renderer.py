from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.schema import ResumePayload

NAME = "LI MI"
LOCATION = "Tucson, AZ"
PHONE = "520 256 8742"
EMAIL = "limi330@yeah.net"
GITHUB_TEXT = "github.com/LiMi686/data-portfolio"
GITHUB_URL = "https://github.com/LiMi686/data-portfolio"

EDUCATION = [
    (
        "University of Arizona",
        "Aug 2024 -- May 2026",
        "Master of Science in Data Science (GPA: 4.0/4.0) | Department Scholarship",
        "Tucson, United States",
    ),
    (
        "Mapua University (Powered by Arizona State U Cintana Alliance Program)",
        "Aug 2019 -- Dec 2023",
        "Bachelor of Science in Data Science (GPA: 3.3/4.0) | Dean's List",
        "Manila, Philippines",
    ),
]

BODY_FONT = "Calibri"
BODY_SIZE = Pt(10.5)
SMALL_SIZE = Pt(10.5)
RIGHT_META_SIZE = Pt(9.5)
NAME_SIZE = Pt(18)
SECTION_SIZE = Pt(14)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT_BLUE = "0563C1"


def _set_run_font(run, *, size=BODY_SIZE, bold=False, italic=False, color: RGBColor | None = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


def _content_width(document: Document):
    section = document.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D3D3D3")
    p_bdr.append(bottom)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT_BLUE)
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), BODY_FONT)
    font.set(qn("w:hAnsi"), BODY_FONT)
    run_props.append(font)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(SMALL_SIZE.pt * 2)))
    run_props.append(size)

    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_header(document: Document) -> None:
    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_after = Pt(2)
    run = name_paragraph.add_run(NAME)
    _set_run_font(run, size=NAME_SIZE, bold=True)

    contact_paragraph = document.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.paragraph_format.space_after = Pt(6)
    run = contact_paragraph.add_run(f"{LOCATION} | {PHONE} | {EMAIL} | ")
    _set_run_font(run, size=SMALL_SIZE)
    _add_hyperlink(contact_paragraph, GITHUB_TEXT, GITHUB_URL)


def _add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title.upper())
    _set_run_font(run, size=SECTION_SIZE, bold=True)
    _add_bottom_border(paragraph)


def _add_two_column_line(
    document: Document,
    left_text: str,
    right_text: str,
    *,
    left_bold: bool = False,
    left_italic: bool = False,
    right_italic: bool = True,
    compact: bool = True,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0 if compact else 1)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.tab_stops.add_tab_stop(_content_width(document), WD_TAB_ALIGNMENT.RIGHT)

    left_run = paragraph.add_run(left_text)
    _set_run_font(left_run, bold=left_bold, italic=left_italic)

    tab_run = paragraph.add_run("\t")
    _set_run_font(tab_run)

    right_run = paragraph.add_run(right_text)
    _set_run_font(right_run, size=RIGHT_META_SIZE, italic=right_italic, color=GREY)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    bullet_run = paragraph.add_run("• ")
    _set_run_font(bullet_run)
    text_run = paragraph.add_run(text)
    _set_run_font(text_run)


def _add_projects(document: Document, payload: ResumePayload) -> None:
    for index, project in enumerate(payload.projects):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        if project.link:
            _add_hyperlink(paragraph, project.title, project.link)
        else:
            title_run = paragraph.add_run(project.title)
            _set_run_font(title_run, bold=True)

        for bullet_text in project.bullets:
            _add_bullet(document, bullet_text)

        if index < len(payload.projects) - 1:
            document.add_paragraph()


def render_cover_letter_docx(output_path: Path, text: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(14)

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    for para_text in paragraphs:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(11)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(para_text)
        run.font.name = BODY_FONT
        run.font.size = Pt(11)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def render_docx(output_path: Path, payload: ResumePayload) -> None:
    document = Document()
    _configure_document(document)
    _add_header(document)

    _add_section_heading(document, "Summary")
    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(2)
    run = summary.add_run(payload.summary)
    _set_run_font(run)

    _add_section_heading(document, "Education")
    for index, (organization, dates, detail, location) in enumerate(EDUCATION):
        _add_two_column_line(document, organization, dates, left_bold=True)
        _add_two_column_line(document, detail, location, left_italic=True)
        if index < len(EDUCATION) - 1:
            document.add_paragraph()

    _add_section_heading(document, "Skills")
    skills = document.add_paragraph()
    skills.paragraph_format.space_after = Pt(2)

    run = skills.add_run("Languages & Analytics: ")
    _set_run_font(run, bold=True)
    run = skills.add_run(payload.skills.languages_analytics)
    _set_run_font(run)

    run = skills.add_run("\nData Stack & ML Frameworks: ")
    _set_run_font(run, bold=True)
    run = skills.add_run(payload.skills.data_stack_ml)
    _set_run_font(run)

    run = skills.add_run("\nDatabases & Concepts: ")
    _set_run_font(run, bold=True)
    run = skills.add_run(payload.skills.databases_concepts)
    _set_run_font(run)

    _add_section_heading(document, "Work & Research Experience")
    for index, experience in enumerate(payload.experiences):
        _add_two_column_line(document, experience.organization, experience.dates, left_bold=True)
        _add_two_column_line(document, experience.role, experience.location, left_italic=True)
        for bullet in experience.bullets:
            _add_bullet(document, bullet)
        if index < len(payload.experiences) - 1:
            document.add_paragraph()

    _add_section_heading(document, "Projects")
    _add_projects(document, payload)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
